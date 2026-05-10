# 檔案 3：docx_generator.py
import os
import pypandoc
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

def clean_latex_spacing(text: str) -> str:
    cleaned = re.sub(r'\$\s+(.*?)\s+\$', r'$\1$', text)
    cleaned = re.sub(r'\$\s+', r'$', cleaned)
    cleaned = re.sub(r'\s+\$', r'$', cleaned)
    return cleaned

def force_kai_font(doc, font_size=12):
    half_pt = str(int(font_size * 2))
    
    for r in doc.element.xpath('.//*[local-name()="r"]'):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
            
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), half_pt)
        
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), half_pt)
        
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), '標楷體')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def generate_word_documents(questions_data: list, template_path: str = None) -> tuple:
    word_md = ""
    download_md = "# 阿凱數學出卷系統 - 測驗卷原始碼\n\n"
    
    for idx, q in enumerate(questions_data, 1):
        clean_text = clean_latex_spacing(q['text'])
        
        # 🚀 判斷是否有圖片與標籤，並將 Markdown 圖片語法塞入指定位置
        img_tag_word = f"\n\n![圖示]({q['img']}){{width=\"3.2in\"}}\n\n" if q.get('img') and os.path.exists(q['img']) else ""
        img_tag_md = f"\n\n![圖示]({q['img']})\n\n" if q.get('img') and os.path.exists(q['img']) else ""

        if "[插入圖片]" in clean_text:
            word_text = clean_text.replace("[插入圖片]", img_tag_word)
            download_text = clean_text.replace("[插入圖片]", img_tag_md)
            
            word_md += f"**{idx}.** {word_text}\n\n<br><br>\n\n"
            download_md += f"### 第 {idx} 題\n\n{download_text}\n\n"
        else:
            # 防呆：如果 AI 忘記加標籤，就維持原本放在最後面的邏輯
            word_md += f"**{idx}.** {clean_text}\n\n"
            download_md += f"### 第 {idx} 題\n\n{clean_text}\n\n"
            
            if q.get('img') and os.path.exists(q['img']):
                word_md += img_tag_word
                download_md += img_tag_md
                
            word_md += "<br><br>\n\n"
            
        if q.get('img') and os.path.exists(q['img']) and q.get('code'):
            download_md += "<details><summary>🖼️ 點擊展開：繪圖 Python 原始碼</summary>\n\n"
            tick3 = chr(96) * 3
            download_md += f"{tick3}python\n" + q['code'] + f"\n{tick3}\n\n</details>\n\n"
            
        download_md += "---\n\n"
        
    temp_teacher_docx = "temp_teacher.docx"
    pypandoc.convert_text(word_md, 'docx', format='md', outputfile=temp_teacher_docx)
    
    output_teacher = "教師用解答本.docx"
    
    if template_path and os.path.exists(template_path):
        master_doc = Document(template_path)
        composer = Composer(master_doc)
        temp_doc = Document(temp_teacher_docx)
        force_kai_font(temp_doc)
        temp_doc.save(temp_teacher_docx)
        
        composer.append(Document(temp_teacher_docx))
        force_kai_font(master_doc)
        master_doc.save(output_teacher)
    else:
        doc = Document(temp_teacher_docx)
        force_kai_font(doc)
        doc.save(output_teacher)
        
    if os.path.exists(temp_teacher_docx):
        os.remove(temp_teacher_docx)

    return output_teacher, download_md